from docx import Document
from docx.shared import Pt
import os

doc = Document()
doc.add_heading('Sign Language Detection System - Code Explanation', 0)

# Intro
doc.add_paragraph('This document outlines the line-by-line architectural decisions, syntax functions, '
                  'and Machine Learning workflows utilized to build a fully automated, real-time Sign Language detection application. '
                  'The project is divided into three distinct modules: Data Collection, Artificial Intelligence Training, and Live Neural Prediction.')

def add_code_explanation(title, explanations):
    doc.add_heading(title, level=1)
    for code, desc in explanations:
        p = doc.add_paragraph()
        if code:
            r1 = p.add_run(code + "\n")
            r1.bold = True
            r1.font.name = 'Consolas'
        r2 = p.add_run(desc)

# --- PART 1: DATA COLLECTION ---
p1_title = "Part 1: The Data Collection Engine"
p1_explanations = [
    ("import cv2, mediapipe as mp, numpy as np, os, time", 
     "Imports OpenCV for video/camera handling, MediaPipe for acquiring biological hand-skeletons, NumPy for matrix/image manipulation, OS for file pathway routing, and time for distinct image saving."),
    ("from mediapipe.tasks import python\nfrom mediapipe.tasks.python import vision", 
     "Imports the advanced Google MediaPipe Tasks API. We specifically target the 'vision' pathway to initialize high-speed skeletal detectors."),
    ("base_options = python.BaseOptions(model_asset_path=r'hand_landmarker.task')\noptions = vision.HandLandmarkerOptions(base_options=base_options, num_hands=1)", 
     "Creates a specific hand-landmark tracker instance, binding it to the locally maintained 'hand_landmarker.task' file. We restrict the tool to detecting a maximum of 1 hand (num_hands=1) to prevent background noise or other people from corrupting our database images."),
    ("HAND_CONNECTIONS = [(0, 1), (1, 2) ...]", 
     "A hardcoded map representing human skeletal constraints (e.g. index finger segments, pinky mappings) ensuring our system knows which of the 21 generated points connect via vectors. This provides the structural skeleton logic."),
    ("os.makedirs(dir_path)", 
     "A directory creation loop checking if 'database/' and 'reference_images/' exist locally. If they do not, it autonomously builds the folders necessary to hold our dataset."),
    ("CANVAS_WIDTH = 1680 \nQUAD_WIDTH = CANVAS_WIDTH // 2", 
     "Defines scaling dimensions for the final output system. 1680x1050 forms the rigid boundary grid for the 4-part multi-module display."),
    ("cap = cv2.VideoCapture(0)", 
     "Initializes binding to webcam index zero (the default primary camera) to poll hardware frames."),
    ("frame = cv2.flip(frame, 1)", 
     "Flips the raw camera data matrix horizontally across the Y-axis. This simulates a 'Mirror' property, making data collection significantly easier for the user intuitively."),
    ("rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)", 
     "Converts the standard OpenCV format (Blue-Green-Red) into traditional Red-Green-Blue matrices, which is the strict, required image channel format expected by the Google MediaPipe algorithm."),
    ("mp_image = mp.Image(image_format=mp.ImageFormat.SRGB, data=rgb_frame)\ndetection_result = detector.detect(mp_image)", 
     "Converts our numpy array matrix into a formalized MediaPipe Image encapsulation class, and pumps it through our loaded ML detector variable to hunt for point-specific structures resulting in a 'detection_result' object."),
    ("black_bg = np.zeros((QUAD_HEIGHT, QUAD_WIDTH, 3), dtype=np.uint8)", 
     "Generates a pitch-black matrix initialized with values of 0. This is the isolated canvas upon which our hand skeleton will be drawn—destroying background noise natively."),
    ("pixels = [(int(lm.x * w), int(lm.y * h)) for lm in hand_landmarks]", 
     "A list comprehension executing structural translation. It extracts standardized float decimals (0 to 1 scaling) from MediaPipe and scales them directly to integers matching our frame pixels to pinpoint the X/Y skeleton."),
    ("cv2.line(...) and cv2.circle(...)", 
     "Uses geometrical vector plotting mathematically mapping circles at the joints and lines between the segments directly onto both the raw webcam map AND our noise-canceling isolated black matrix."),
    ("cropped_img_to_save = black_bg[y_min:y_max, x_min:x_max]", 
     "Analyzes the boundaries (max/min X and Y), adds spatial padding logic of 50px, and isolates solely the skeleton array to cut away useless empty space."),
    ("img_resize = cv2.resize(cropped_img_to_save, (w_cal, img_size))", 
     "Scales the bounded skeletal image precisely into the center of a newly allocated 300x300 pitch black canvas without destroying the aspect ratio geometry of the captured sign."),
    ("cv2.imwrite(img_name, img_bg)", 
     "Performs the physical write onto the physical hard drive, caching away training data structurally segregated into alphabet-named subfolders over time.")
]
add_code_explanation(p1_title, p1_explanations)

# --- PART 2: NEURAL NETWORK ---
p2_title = "Part 2: TensorFlow Neural Network Generation"
p2_explanations = [
    ("import tensorflow as tf\nfrom tensorflow.keras.models import Sequential", 
     "Draws upon the Keras deep learning framework from Google's TensorFlow suite. 'Sequential' explicitly designates that we are building a linear topological stack of neural layers flowing consecutively."),
    ("img_size = 100", 
     "Forces a normalization standard. Though our generator crafted 300x300 matrices, transforming them down to 100x100 compresses processing payload dramatically without drastically altering the shape variables required by the algorithm."),
    ("img_arr = cv2.imread(..., cv2.IMREAD_GRAYSCALE)", 
     "Compresses mathematical parameters significantly. By crushing RGB color mapping into singular Grayscale tensors, our Convolutional algorithm avoids memorizing skin tone or room color, locking focus fully onto the contrasting structural lines."),
    ("X = np.array(X).reshape(-1, img_size, img_size, 1)", 
     "Transforms the unstructured Python list 'X' into a concrete Numpy block architecture spanning [Samples count, 100 Height, 100 Width, 1 Channel]."),
    ("X = X / 255.0", 
     "Normalizes our internal tensor logic. Division of color (255) maps elements between 0.0 and 1.0. This prevents massive scalar weights from destructively overpowering activation function boundaries in neural nodes."),
    ("y = to_categorical(np.array(y), num_classes=num_classes)", 
     "One-hot encodes scalar category labels via positional matrices (e.g. integer '3' mapping to standard [0,0,0,1,...]) guaranteeing our final layer correctly identifies probabilistic categorical outputs natively."),
    ("X_train, X_test... = train_test_split(..., test_size=0.2)", 
     "Enforces ML standards by hiding 20% of the dataset as 'test' references to evaluate the model without cheating by testing strictly unseen skeleton structures to guarantee general real-world applicability."),
    ("Conv2D(32, (3, 3), activation='relu')", 
     "Constructs Convolutional extraction filters searching across a 3x3 topological pool for skeletal features (arcs, lines, spatial density). Utilizing 'ReLU' (Rectified Linear Unit) activation removes negative matrices and solves gradient extinction logic mathematically."),
    ("MaxPooling2D((2, 2))", 
     "Down-samples data maps natively. The strongest structural features remain internally bound while dropping generalized spatial filler via a 2x2 local maximum filter."),
    ("Flatten() \nDense(128, activation='relu')", 
     "Smashes multiple 3rd-dimensional filtered node systems down into a stark 1D flat dense network of fully connected logic nodes passing high-tier categorical information representations."),
    ("Dropout(0.5)", 
     "Severely blocks overfitting. This aggressively ignores/kills 50% of randomized artificial weights every cycle, guaranteeing the neural network finds robust internal pathways mapping signs rather than heavily biasing upon specific nodes mimicking test data."),
    ("Dense(num_classes, activation='softmax')", 
     "Our terminal layer. Uses Softmax mathematical probability smoothing, providing standard values denoting percentage of certainty dynamically balanced across mapping our detected alphabet values structurally."),
    ("model.compile(optimizer='adam', loss='categorical_crossentropy')", 
     "Attaches the 'Adam' optimizer to direct layer convergence using categorical cross_entropy targeting exact loss differentials to adjust mapping filters automatically."),
    ("model.save(r'sign_model.h5')", 
     "Serializes the compiled internal neural structure, freezing the trained numerical weights definitively into an '.h5' file layout mapped onto the physical drive.")
]
add_code_explanation(p2_title, p2_explanations)

# --- PART 3: LIVE PREDICTION ---
p3_title = "Part 3: Live Prediction & Real-Time Output Demo"
p3_explanations = [
    ("model = tf.keras.models.load_model(r'd:\Image Project\sign_model.h5')", 
     "De-serializes our internal, completely trained neural weights architecture mapping the structure completely back into RAM for mathematical propagation natively."),
    ("labels.append(line.strip().split(' ')[1])", 
     "Parses our generated labels dictionary sequentially matching final layer node indexes precisely back into human-readable characters natively output as lists."),
    ("nn_input = nn_input.reshape(-1, 100, 100, 1) / 255.0", 
     "Mirrors our training sequence normalization entirely, applying an indistinguishable scale format scaling to 1.0 against the actively processed webcam skeletal stream bounding boxes ensuring logic inputs overlap mapping expectations."),
    ("prediction = model.predict(nn_input, verbose=0)", 
     "A single pass injection into the deep learning pipeline dynamically outputting an array representing internal prediction bounds correlating certainty per class index strictly mathematically."),
    ("class_idx = np.argmax(prediction[0])", 
     "Executes indexing arrays mapped locally against the maximum internal classification rating logically, locating the specific alphabet index directly representing highest statistical probability."),
    ("cv2.putText(output_bg, f'Prediction: {predicted_character}', ...)", 
     "Reimplements our live certainty bindings onto a dynamically updating output UI. The detected predictions execute visual rendering instantly providing user visual feedback validation."),
    ("word_accumulator += predicted_character", 
     "Acts via discrete listener bindings tied directly into Python arrays, compiling strings structurally forming word aggregations conditionally when activated manually via key mapping.")
]
add_code_explanation(p3_title, p3_explanations)

# Finalization
doc.save(r'd:\Image Project\Project_Submission_Code_Explanation.docx')
